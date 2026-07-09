from __future__ import annotations

import io
import mimetypes
import re
from dataclasses import dataclass
from pathlib import Path

from backend.utils.settings import get_settings

settings = get_settings()


@dataclass(slots=True)
class PreparedKnowledgeDocument:
    filename: str
    category: str
    extracted_text: str
    size_label: str
    content_type: str
    raw_bytes: bytes


@dataclass(slots=True)
class DownloadArtifact:
    filename: str
    content_type: str
    content: bytes


class KnowledgeIngestionService:
    def __init__(self) -> None:
        self._upload_dir = Path(settings.get_uploads_path())
        self._upload_dir.mkdir(parents=True, exist_ok=True)

    def infer_category(self, filename: str, provided_category: str | None = None) -> str:
        if provided_category and provided_category.strip():
            return provided_category.strip()

        filename_lower = filename.lower()
        if "contract" in filename_lower:
            return "Supplier Contracts"
        if "invoice" in filename_lower:
            return "Invoice"
        if "catalog" in filename_lower or "catalogue" in filename_lower:
            return "Medicine Catalogue"
        if "policy" in filename_lower:
            return "Procurement Policy"
        return "Hospital SOP"

    def prepare_text_document(
        self,
        *,
        filename: str,
        category: str,
        content: str,
        size_label: str | None = None,
    ) -> PreparedKnowledgeDocument:
        raw_bytes = content.encode("utf-8")
        extracted_text = self._normalize_text(content)
        if len(extracted_text) < 20:
            raise ValueError("Knowledge content is too short to index.")

        return PreparedKnowledgeDocument(
            filename=filename.strip() or "knowledge-note.txt",
            category=self.infer_category(filename, category),
            extracted_text=extracted_text,
            size_label=size_label or self.build_size_label(len(raw_bytes)),
            content_type=self._guess_content_type(filename, "text/plain"),
            raw_bytes=raw_bytes,
        )

    def prepare_binary_document(
        self,
        *,
        filename: str,
        category: str | None,
        raw_bytes: bytes,
        content_type: str | None = None,
    ) -> PreparedKnowledgeDocument:
        safe_filename = filename.strip()
        if not safe_filename:
            raise ValueError("Uploaded files must include a filename.")
        if not raw_bytes:
            raise ValueError("Uploaded files cannot be empty.")

        extracted_text = self._normalize_text(self._extract_text(safe_filename, raw_bytes, content_type))
        if len(extracted_text) < 20:
            raise ValueError("The uploaded file did not contain enough readable text to index.")

        return PreparedKnowledgeDocument(
            filename=safe_filename,
            category=self.infer_category(safe_filename, category),
            extracted_text=extracted_text,
            size_label=self.build_size_label(len(raw_bytes)),
            content_type=self._guess_content_type(safe_filename, content_type),
            raw_bytes=raw_bytes,
        )

    def build_summary(self, *, category: str, extracted_text: str, chunk_count: int) -> str:
        preview = next((part.strip() for part in re.split(r"[.\n]", extracted_text) if part.strip()), "")
        preview = preview[:140].rstrip()
        summary = f"{category} indexed into {chunk_count} retrieval chunks for grounded hospital operations."
        if preview:
            summary = f"{summary} {preview}."
        return summary

    def persist_original(self, file_code: str, filename: str, raw_bytes: bytes) -> Path:
        self.delete_original(file_code)
        target = self._upload_dir / self._storage_name(file_code, filename)
        target.write_bytes(raw_bytes)
        return target

    def delete_original(self, file_code: str) -> None:
        for path in self._upload_dir.glob(f"{file_code}__*"):
            path.unlink(missing_ok=True)

    def build_download_artifact(self, file_code: str, filename: str, fallback_text: str) -> DownloadArtifact:
        stored_asset = next(self._upload_dir.glob(f"{file_code}__*"), None)
        if stored_asset is not None and stored_asset.exists():
            return DownloadArtifact(
                filename=filename,
                content_type=self._guess_content_type(filename, None),
                content=stored_asset.read_bytes(),
            )

        fallback_name = f"{Path(filename).stem or file_code}.txt"
        return DownloadArtifact(
            filename=fallback_name,
            content_type="text/plain; charset=utf-8",
            content=fallback_text.encode("utf-8"),
        )

    def build_size_label(self, byte_count: int) -> str:
        if byte_count < 1024:
            return f"{byte_count} B"
        if byte_count < 1024 * 1024:
            return f"{byte_count / 1024:.1f} KB"
        return f"{byte_count / 1024 / 1024:.1f} MB"

    def _extract_text(self, filename: str, raw_bytes: bytes, content_type: str | None) -> str:
        suffix = Path(filename).suffix.lower()
        text_like_suffixes = {".txt", ".md", ".csv", ".json", ".log", ".xml"}

        if suffix in text_like_suffixes or (content_type or "").startswith("text/"):
            return raw_bytes.decode("utf-8", errors="ignore")

        if suffix == ".pdf":
            try:
                from pypdf import PdfReader
            except ImportError as exc:  # pragma: no cover - import protection
                raise ValueError("PDF ingestion is unavailable because pypdf is not installed.") from exc

            reader = PdfReader(io.BytesIO(raw_bytes))
            pages = []
            for page in reader.pages:
                text = (page.extract_text() or "").strip()
                if text:
                    pages.append(text)
            return "\n\n".join(pages)

        if suffix == ".docx":
            try:
                from docx import Document
            except ImportError as exc:  # pragma: no cover - import protection
                raise ValueError("DOCX ingestion is unavailable because python-docx is not installed.") from exc

            document = Document(io.BytesIO(raw_bytes))
            return "\n".join(paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip())

        raise ValueError("Supported knowledge file types are TXT, PDF, and DOCX.")

    def _normalize_text(self, text: str) -> str:
        normalized = text.replace("\r\n", "\n").replace("\r", "\n")
        normalized = re.sub(r"\n{3,}", "\n\n", normalized)
        normalized = re.sub(r"[ \t]{2,}", " ", normalized)
        return normalized.strip()

    def _storage_name(self, file_code: str, filename: str) -> str:
        extension = Path(filename).suffix or ".bin"
        stem = re.sub(r"[^A-Za-z0-9._-]+", "-", Path(filename).stem).strip("-._") or "document"
        return f"{file_code}__{stem}{extension.lower()}"

    def _guess_content_type(self, filename: str, fallback: str | None) -> str:
        guessed = mimetypes.guess_type(filename)[0]
        return guessed or fallback or "application/octet-stream"


knowledge_ingestion_service = KnowledgeIngestionService()
